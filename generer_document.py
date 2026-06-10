"""
Génère le document Word de présentation de SharCode.
Lancer avec : python generer_document.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Couleurs ──────────────────────────────────────────────────
BLEU_SHARK  = RGBColor(0x00, 0x77, 0xB6)
BLEU_FONCE  = RGBColor(0x03, 0x04, 0x5E)
GRIS        = RGBColor(0x64, 0x74, 0x8B)
NOIR        = RGBColor(0x1E, 0x29, 0x3B)
BLANC       = RGBColor(0xFF, 0xFF, 0xFF)


def creer_document():
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    # ── Styles de base ────────────────────────────────────────
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = NOIR
    style_normal.paragraph_format.space_after = Pt(8)
    style_normal.paragraph_format.line_spacing = Pt(16)

    # ─────────────────────────────────────────────────────────
    #  PAGE DE TITRE
    # ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.add_run('🦈  SharCode')
    run.font.size   = Pt(36)
    run.font.bold   = True
    run.font.color.rgb = BLEU_SHARK

    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sous_titre.add_run('Un langage de programmation en français naturel')
    run.font.size   = Pt(16)
    run.font.color.rgb = BLEU_FONCE
    run.font.italic = True

    doc.add_paragraph()

    ligne = doc.add_paragraph()
    ligne.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ligne.add_run('Projet Annuel — 2025 / 2026')
    run.font.size  = Pt(12)
    run.font.color.rgb = GRIS

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────
    #  SECTION 1 — L'ORIGINE DU PROJET
    # ─────────────────────────────────────────────────────────
    ajouter_titre1(doc, '1. L\'origine du projet')

    ajouter_paragraphe(doc,
        "Apprendre à programmer est aujourd'hui présenté comme une compétence "
        "essentielle. Les collèges intègrent Scratch dans leurs cours de technologie, "
        "et Python arrive progressivement au lycée. Mais entre les deux, il existe "
        "un vide que personne n'a encore vraiment comblé."
    )

    ajouter_paragraphe(doc,
        "Scratch, c'est de la programmation par blocs visuels. C'est accessible, "
        "ludique, et efficace pour comprendre qu'un programme est une suite "
        "d'instructions. Mais dès qu'un élève passe à un vrai langage texte comme "
        "Python, le choc est souvent brutal :"
    )

    ajouter_liste(doc, [
        "La syntaxe est entièrement en anglais — if, while, def, return, True, False…",
        "Les messages d'erreur sont techniques et en anglais",
        "Les règles d'indentation sont invisibles et provoquent des erreurs incompréhensibles",
        "La documentation, les tutoriels, les forums — tout est en anglais",
    ])

    ajouter_paragraphe(doc,
        "Pour les 320 millions de francophones dans le monde, et en particulier "
        "pour les collégiens qui ne maîtrisent pas encore l'anglais technique, "
        "cette double barrière — apprendre la logique ET le faire en anglais — "
        "est un frein réel. Des travaux en sciences de l'éducation montrent qu'un "
        "apprenant comprend mieux un concept nouveau quand il est présenté dans sa "
        "langue maternelle. Quand on n'a pas à traduire dans sa tête, on peut "
        "vraiment réfléchir à la logique."
    )

    ajouter_paragraphe(doc,
        "C'est de cette observation qu'est né SharCode."
    )

    # ─────────────────────────────────────────────────────────
    #  SECTION 2 — PRÉSENTATION DE SHARCODE
    # ─────────────────────────────────────────────────────────
    ajouter_titre1(doc, '2. Présentation de SharCode')

    ajouter_titre2(doc, 'Un langage qui se lit comme du français')

    ajouter_paragraphe(doc,
        "SharCode est un langage de programmation dont la syntaxe est écrite en "
        "français naturel. L'idée centrale est simple : retirer la barrière de la "
        "langue pour que l'élève puisse se concentrer sur ce qui compte — la logique."
    )

    ajouter_paragraphe(doc,
        "Voici un exemple de programme SharCode :"
    )

    ajouter_code(doc,
        'age vaut 17\n\n'
        'si age >= 18 alors\n'
        '    ecrire "Tu es majeur."\n'
        'sinon\n'
        '    ecrire "Tu seras majeur dans " + (18 - age) + " an(s)."\n'
        'fin'
    )

    ajouter_paragraphe(doc,
        "Un élève de 4ème qui n'a jamais écrit une ligne de code comprend ce "
        "programme en le lisant à voix haute. Il n'y a rien à décoder, rien à "
        "traduire. La logique apparaît directement."
    )

    ajouter_titre2(doc, 'Une ressemblance avec Python assumée')

    ajouter_paragraphe(doc,
        "SharCode ressemble à Python. C'est voulu, et nous l'assumons complètement. "
        "Le but du projet n'est pas d'inventer un langage radicalement nouveau — "
        "c'est de créer le meilleur chemin possible vers Python pour un élève francophone."
    )

    ajouter_paragraphe(doc,
        "Voici la même condition écrite dans les deux langages, côte à côte :"
    )

    ajouter_comparaison(doc,
        "SharCode",
        'si note >= 10 alors\n    ecrire "Admis"\nsinon\n    ecrire "Recalé"\nfin',
        "Python",
        'if note >= 10:\n    print("Admis")\nelse:\n    print("Recalé")'
    )

    ajouter_paragraphe(doc,
        "La structure est identique. La logique est la même. Mais en SharCode, "
        "un collégien francophone peut lire et écrire ce code sans jamais avoir "
        "appris un mot d'anglais. C'est précisément là que réside la valeur "
        "ajoutée du projet."
    )

    ajouter_titre2(doc, 'Ce qui différencie SharCode de Python en pratique')

    ajouter_paragraphe(doc,
        "Au-delà des mots-clés français, SharCode introduit plusieurs simplifications "
        "qui réduisent les erreurs fréquentes chez les débutants :"
    )

    ajouter_liste(doc, [
        "Les blocs se ferment avec 'fin' — il n'y a pas d'erreurs d'indentation "
        "invisibles, qui sont la première source de blocage en Python.",

        "'vaut' pour assigner une valeur, '==' pour comparer — la confusion entre "
        "= et == qui fait trébucher tous les débutants disparaît complètement.",

        "La boucle 'pour i de 1 a 10' est immédiatement lisible, là où "
        "'for i in range(1, 11)' demande une explication supplémentaire.",

        "Les messages d'erreur sont rédigés en français, avec des explications "
        "adaptées à un débutant : \"La variable 'x' n'existe pas. As-tu pensé "
        "à l'initialiser ?\" plutôt que \"NameError: name 'x' is not defined\".",

        "Aucun code superflu : pas d'import, pas de if __name__ == '__main__', "
        "pas de gestion de classes obligatoire. On écrit, ça tourne.",
    ])

    ajouter_paragraphe(doc,
        "Un élève qui maîtrise SharCode connaît déjà toute la logique de la "
        "programmation. Passer à Python ne demande alors qu'une chose : apprendre "
        "les mots anglais correspondants. Ce qui peut prendre des mois en partant "
        "de zéro prend quelques semaines quand la logique est déjà acquise."
    )

    # ─────────────────────────────────────────────────────────
    #  SECTION 3 — INTÉGRATION SCOLAIRE
    # ─────────────────────────────────────────────────────────
    ajouter_titre1(doc, '3. L\'intégration dans le parcours scolaire')

    ajouter_titre2(doc, 'Un chaînon manquant dans les programmes')

    ajouter_paragraphe(doc,
        "En France, le programme officiel de technologie au collège intègre Scratch "
        "dès la 6ème pour initier les élèves à la pensée algorithmique. Python "
        "apparaît ensuite au lycée dans les cours de mathématiques et de NSI. "
        "Mais il n'existe rien entre les deux pour accompagner la transition."
    )

    ajouter_schema(doc)

    ajouter_paragraphe(doc,
        "SharCode est conçu pour occuper précisément cet espace : en 4ème et 3ème, "
        "quand les élèves ont déjà les bases conceptuelles apportées par Scratch, "
        "mais ne sont pas encore prêts pour la rigueur syntaxique de Python en anglais."
    )

    ajouter_titre2(doc, 'Comment SharCode s\'utilise en classe')

    ajouter_paragraphe(doc,
        "SharCode fonctionne directement dans le navigateur web. Aucune installation "
        "n'est nécessaire sur les postes des élèves — ce qui est un avantage majeur "
        "dans un contexte scolaire où les droits d'administration sont souvent "
        "restreints. Un enseignant partage un lien, les élèves ouvrent leur navigateur "
        "et commencent à coder immédiatement."
    )

    ajouter_paragraphe(doc,
        "L'interface intègre directement :"
    )

    ajouter_liste(doc, [
        "Un éditeur de code avec numérotation des lignes",
        "6 leçons progressives, de la variable à la fonction récursive, "
        "accessibles en un clic depuis la barre latérale",
        "Une console de sortie qui affiche les résultats du programme",
        "Des exemples prêts à l'emploi pour explorer et modifier",
        "Une référence rapide de toute la syntaxe",
    ])

    ajouter_titre2(doc, 'A qui s\'adresse SharCode')

    ajouter_paragraphe(doc,
        "La cible principale est constituée des élèves de 4ème et 3ème. Mais "
        "SharCode s'adresse également à :"
    )

    ajouter_liste(doc, [
        "Les enseignants de technologie qui souhaitent aller au-delà de Scratch "
        "sans passer directement à Python",
        "Les adultes en reconversion professionnelle qui veulent découvrir la "
        "programmation sans la barrière de l'anglais",
        "Plus largement, les 320 millions de francophones dans le monde, "
        "notamment en Afrique subsaharienne et au Maghreb, pour qui les outils "
        "d'apprentissage du code sont quasi exclusivement en anglais",
    ])

    # ─────────────────────────────────────────────────────────
    #  SECTION 4 — CE QU'ON A CONSTRUIT
    # ─────────────────────────────────────────────────────────
    ajouter_titre1(doc, '4. Ce que nous avons construit')

    ajouter_paragraphe(doc,
        "SharCode est un projet complet, fonctionnel, développé de zéro. "
        "Il comprend plusieurs composants :"
    )

    ajouter_titre2(doc, 'L\'interpréteur')

    ajouter_paragraphe(doc,
        "Le cœur du projet est un interpréteur de langage écrit en Python, "
        "composé de quatre parties distinctes :"
    )

    ajouter_liste(doc, [
        "Le lexer (analyseur lexical) : lit le code source mot par mot et "
        "produit une liste de tokens — les unités de base du langage. Il gère "
        "notamment les mots-clés composés ('tant que', 'sinon si') et accepte "
        "indifféremment les accents ('écrire' et 'ecrire' fonctionnent tous les deux).",

        "Le parser (analyseur syntaxique) : vérifie que la structure du code "
        "est correcte et construit un Arbre Syntaxique Abstrait (AST) — une "
        "représentation en mémoire du programme.",

        "L'interpréteur : parcourt l'arbre et exécute les instructions une à une, "
        "en gérant les variables, les portées, les appels de fonctions et la récursivité.",

        "Le gestionnaire d'erreurs : intercepte tous les problèmes et produit des "
        "messages clairs en français, avec le numéro de ligne concerné.",
    ])

    ajouter_titre2(doc, 'L\'interface web')

    ajouter_paragraphe(doc,
        "Une application web développée avec Flask (Python) pour le serveur et "
        "HTML / CSS / JavaScript pour l'interface. Elle permet d'écrire et "
        "d'exécuter du code SharCode directement dans le navigateur, sans rien "
        "installer. L'interface comprend un éditeur, une console, des leçons "
        "intégrées et des exemples."
    )

    ajouter_titre2(doc, 'La documentation')

    ajouter_paragraphe(doc,
        "Une documentation complète du langage a été rédigée : référence de toute "
        "la syntaxe, guide de démarrage, exemples commentés, et explication du "
        "positionnement pédagogique du projet."
    )

    # ─────────────────────────────────────────────────────────
    #  SECTION 5 — CONCLUSION
    # ─────────────────────────────────────────────────────────
    ajouter_titre1(doc, '5. Conclusion')

    ajouter_paragraphe(doc,
        "SharCode répond à une question simple : pourquoi est-ce qu'apprendre à "
        "coder en France demande de savoir lire l'anglais ?"
    )

    ajouter_paragraphe(doc,
        "Nous ne prétendons pas avoir inventé un nouveau paradigme de programmation. "
        "Nous avons construit un outil pratique, pour un contexte précis, qui "
        "comble un vide réel. SharCode ressemble à Python parce que l'objectif "
        "est que les élèves arrivent en Python avec de l'avance — pas qu'ils "
        "restent sur SharCode. C'est un tremplin, pas une destination."
    )

    ajouter_paragraphe(doc,
        "La valeur du projet est là : dans sa clarté, dans son honnêteté sur "
        "ce qu'il est, et dans la réalité du besoin auquel il répond."
    )

    citation = doc.add_paragraph()
    citation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    citation.paragraph_format.space_before = Pt(20)
    citation.paragraph_format.space_after  = Pt(20)
    run = citation.add_run('"Coder comme on parle."')
    run.font.size   = Pt(14)
    run.font.italic = True
    run.font.color.rgb = BLEU_SHARK

    doc.save('SharCode_Presentation.docx')
    print("Document genere : SharCode_Presentation.docx")


# ── Fonctions utilitaires ─────────────────────────────────────

def ajouter_titre1(doc, texte):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(texte)
    run.font.size  = Pt(16)
    run.font.bold  = True
    run.font.color.rgb = BLEU_FONCE
    # Ligne de séparation sous le titre
    ajouter_bordure_bas(p)


def ajouter_titre2(doc, texte):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(texte)
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = BLEU_SHARK


def ajouter_paragraphe(doc, texte):
    p = doc.add_paragraph(texte)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = NOIR


def ajouter_liste(doc, elements):
    for el in elements:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent  = Cm(0.75)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.line_spacing = Pt(15)
        run = p.add_run(el)
        run.font.size = Pt(11)
        run.font.color.rgb = NOIR


def ajouter_code(doc, texte):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    # Fond gris via ombrage
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A1B2E')
    pPr.append(shd)
    run = p.add_run(texte)
    run.font.name  = 'Consolas'
    run.font.size  = Pt(10)
    run.font.color.rgb = RGBColor(0xCD, 0xD6, 0xF4)


def ajouter_comparaison(doc, titre_g, code_g, titre_d, code_d):
    """Affiche deux blocs de code côte à côte (simulation avec tableau)."""
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    # Titres
    for i, titre in enumerate([titre_g, titre_d]):
        cell = table.cell(0, i)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(titre)
        run.font.bold  = True
        run.font.size  = Pt(10)
        run.font.color.rgb = BLANC
        # Fond coloré
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '0077B6' if i == 0 else '03045E')
        tcPr.append(shd)

    # Code
    couleurs = ['1A1B2E', '0D1117']
    for i, code in enumerate([code_g, code_d]):
        cell = table.cell(1, i)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(code)
        run.font.name  = 'Consolas'
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0xCD, 0xD6, 0xF4)
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  couleurs[i])
        tcPr.append(shd)

    doc.add_paragraph()


def ajouter_schema(doc):
    """Schéma de progression textuel."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)

    blocs = [
        ('6ème / 5ème', 'Scratch', '(blocs visuels)'),
        ('4ème / 3ème', 'SharCode', '(texte, français)'),
        ('Lycée',       'Python / JS', '(texte, anglais)'),
    ]

    for i, (niveau, nom, desc) in enumerate(blocs):
        couleur = ['0077B6', '00B4D8', '03045E'][i]
        run = p.add_run(f'  {niveau}\n{nom}\n{desc}  ')
        run.font.name  = 'Calibri'
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(
            int(couleur[:2], 16),
            int(couleur[2:4], 16),
            int(couleur[4:], 16)
        )
        run.font.bold = (nom == 'SharCode')
        if i < 2:
            p.add_run('  →  ').font.size = Pt(12)

    doc.add_paragraph()


def ajouter_bordure_bas(paragraphe):
    """Ajoute une fine ligne sous un paragraphe (style titre)."""
    pPr = paragraphe._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0077B6')
    pBdr.append(bottom)
    pPr.append(pBdr)


if __name__ == '__main__':
    creer_document()
